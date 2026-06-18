from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Dokumentimi_Projekti_Ecommerce_Librave.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(88, 98, 110)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
BORDER = "C9D3DF"


def font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_style(cell, fill=None):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", 90), ("start", 130), ("bottom", 90), ("end", 130)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        item = borders.find(qn(tag))
        if item is None:
            item = OxmlElement(tag)
            borders.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), "6")
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), BORDER)

    if fill:
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_style(cell)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Title", 24, RGBColor(0, 0, 0), 0, 4),
        ("Subtitle", 14, MUTED, 0, 16),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name.startswith("Heading")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.text = "Platforme e-commerce per shitjen e librave"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(header.runs[0], 9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Dokumentim i sherbimeve te platformes"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.runs[0], 9, color=MUTED)


def para(doc, text="", style=None, size=11, bold=False, color=INK, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    for r in p.runs:
        font(r, size=16 if level == 1 else 13, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        font(r)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        font(r)


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_style(cell, CALLOUT_FILL)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    font(r2, size=10.5)
    para(doc, "", after=4)


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    for i, header in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        set_cell_style(c, LIGHT_FILL)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        font(r, size=10.5, bold=True)
    for row_data in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.10
            r = p.add_run(value)
            font(r, size=10)
    set_table_geometry(t, widths)
    para(doc, "", after=5)


def title_page(doc):
    para(doc, "DOKUMENTIM PROJEKTI", size=10.5, bold=True, color=MUTED, after=6)
    title = doc.add_paragraph(style="Title")
    title.add_run("Platforme E-commerce per Shitjen e Librave")
    for r in title.runs:
        font(r, 24, bold=True, color=RGBColor(0, 0, 0))
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Pershkrim i platformes, sherbimeve dhe funksioneve kryesore")
    for r in subtitle.runs:
        font(r, 14, color=MUTED)

    rows = [
        ("Emri i projektit", "Book Management System"),
        ("Lloji i platformes", "Dyqan online per libra"),
        ("Qellimi", "Shfletim, perzgjedhje, vleresim dhe porositje librash online"),
        ("Perdoruesit kryesore", "Kliente, vizitore dhe administrator/e te platformes"),
        ("Data", "22 maj 2026"),
    ]
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        font(r, bold=True)
        rv = p.add_run(value)
        font(rv)

    callout(
        doc,
        "Fokusi i dokumentit",
        "Ky dokument pershkruan se cfare ofron platforma per perdoruesit dhe per menaxhimin e nje dyqani "
        "online librash. Nuk trajton kodin, strukturen teknike apo detaje te implementimit.",
    )


def build():
    doc = Document()
    configure(doc)
    title_page(doc)

    heading(doc, "1. Permbledhje e platformes")
    para(
        doc,
        "Platforma e-commerce per shitjen e librave eshte nje zgjidhje online qe u mundeson perdoruesve "
        "te shfletojne katalogun e librave, te lexojne informacione per secilin liber, te krijojne llogari, "
        "te ruajne libra te preferuar, te japin vleresime, te komentojne dhe te krijojne porosi. Platforma "
        "sherben si ure lidhese mes lexuesve dhe dyqanit, duke e bere procesin e blerjes me te shpejte, "
        "me te organizuar dhe me te qasshem nga distanca.",
    )

    heading(doc, "2. Qellimi i platformes")
    bullets(
        doc,
        [
            "Te ofroje nje menyre te thjeshte per gjetjen dhe blerjen e librave online.",
            "Te ndihmoje dyqanin ne prezantimin e librave sipas autoreve dhe kategorive.",
            "Te krijoje eksperience me personale per perdoruesit permes favoriteve, komenteve dhe vleresimeve.",
            "Te lehtesoje menaxhimin e porosive dhe ndjekjen e statusit te tyre.",
            "Te rrise dukshmerine e librave dhe te permiresoje komunikimin me klientet.",
        ],
    )

    heading(doc, "3. Perdoruesit e platformes")
    table(
        doc,
        ["Perdoruesi", "Cfare mund te beje", "Vlera qe merr"],
        [
            ("Vizitori", "Shfleton libra, autore dhe kategori pa pasur domosdoshmerisht llogari.", "Njihet me oferten e dyqanit dhe gjen libra me interes."),
            ("Klienti i regjistruar", "Identifikohet, shton libra ne favorite, komenton, vlereson dhe krijon porosi.", "Ka eksperience personale dhe ruan historikun/zgjedhjet e veta."),
            ("Administratori", "Menaxhon librat, autoret, kategorite dhe statusin e porosive.", "Mban katalogun dhe shitjet te organizuara."),
        ],
        [1800, 4350, 3210],
    )

    heading(doc, "4. Sherbimet kryesore qe ofron platforma")
    heading(doc, "4.1 Katalogu online i librave", level=2)
    para(
        doc,
        "Platforma ofron nje katalog digjital ku paraqiten librat qe jane ne shitje. Cdo liber mund te "
        "shfaqe titullin, pershkrimin, cmimin, vitin e publikimit, autorin, kategorine dhe imazhin. Ky "
        "sherbim i ndihmon klientet te krahasojne libra dhe te marrin vendim me te informuar para blerjes.",
    )
    heading(doc, "4.2 Organizimi sipas autoreve dhe kategorive", level=2)
    para(
        doc,
        "Librat organizohen sipas autoreve dhe kategorive, duke e bere shfletimin me te lehte. Perdoruesi "
        "mund te gjeje libra nga nje autor i caktuar ose nga nje fushe qe i intereson, si romane, shkence, "
        "histori, biznes, teknologji ose kategori te tjera te dyqanit.",
    )
    heading(doc, "4.3 Regjistrimi dhe hyrja ne llogari", level=2)
    para(
        doc,
        "Platforma u mundeson klienteve te krijojne llogari personale dhe te hyjne ne sistem. Llogaria e "
        "perdoruesit eshte e rendesishme per sherbime si porosite personale, lista e favoriteve, komentet "
        "dhe vleresimet.",
    )
    heading(doc, "4.4 Lista e librave te preferuar", level=2)
    para(
        doc,
        "Klienti mund te ruaje librat qe i pelqejne ne listen e favoriteve. Ky sherbim eshte i dobishem "
        "kur perdoruesi deshiron te kthehet me vone te nje liber, ta krahasoje me libra te tjere ose ta "
        "bleje ne nje moment tjeter.",
    )
    heading(doc, "4.5 Komente dhe vleresime", level=2)
    para(
        doc,
        "Perdoruesit mund te komentojne dhe te vleresojne librat me yje. Komentet dhe vleresimet ndihmojne "
        "bleresit e tjere te kuptojne me mire cilesine, permbajtjen dhe pershtatshmerine e nje libri.",
    )
    heading(doc, "4.6 Porositja e librave", level=2)
    para(
        doc,
        "Platforma mundeson krijimin e porosive nga librat e zgjedhur. Gjate porosise llogaritet cmimi "
        "total dhe ruhet statusi i porosise. Kjo e ben procesin e blerjes te strukturuar dhe te gjurmueshem.",
    )

    heading(doc, "5. Sherbimet per klientin")
    table(
        doc,
        ["Sherbimi", "Pershkrimi"],
        [
            ("Shfletim i katalogut", "Klienti sheh listen e librave te disponueshem dhe detajet e secilit liber."),
            ("Informacion per librin", "Platforma paraqet pershkrimin, cmimin, vitin, autorin, kategorine dhe imazhin."),
            ("Llogari personale", "Klienti regjistrohet dhe perdor sherbime qe lidhen me profilin e tij."),
            ("Favorite", "Klienti ruan libra per t'i pare ose blere me vone."),
            ("Vleresime", "Klienti jep yje per librat dhe sheh vleresimin mesatar."),
            ("Komente", "Klienti ndan mendimin e tij dhe lexon pervojen e perdoruesve te tjere."),
            ("Porosi", "Klienti zgjedh librat dhe krijon porosi online."),
        ],
        [2450, 6910],
    )

    heading(doc, "6. Sherbimet per administratorin")
    bullets(
        doc,
        [
            "Shtimi i librave te rinj ne katalog me te dhena pershkruese dhe imazh.",
            "Perditesimi i cmimeve, pershkrimeve, autoreve dhe kategorive.",
            "Fshirja e librave qe nuk jane me pjese e ofertes.",
            "Menaxhimi i autoreve dhe kategorive per ta mbajtur katalogun te organizuar.",
            "Ndjekja dhe perditesimi i statusit te porosive.",
            "Monitorimi i komenteve dhe vleresimeve per te kuptuar interesin e klienteve.",
        ],
    )

    heading(doc, "7. Rrjedha e perdorimit nga klienti")
    numbered(
        doc,
        [
            "Klienti hap platformen dhe shfleton katalogun e librave.",
            "Zgjedh nje liber dhe lexon detajet e tij.",
            "Nese deshiron, krijon llogari ose hyn ne llogarine ekzistuese.",
            "Shton libra ne favorite ose vazhdon drejt porosise.",
            "Krijon porosine me librat e zgjedhur.",
            "Pas blerjes mund te komentoje ose te vleresoje librin.",
        ],
    )

    heading(doc, "8. Perfitimet e platformes")
    table(
        doc,
        ["Perfitimi", "Shpjegimi"],
        [
            ("Qasje 24/7", "Klientet mund te shfletojne dhe te porosisin libra ne cdo kohe."),
            ("Eksperience me e shpejte", "Informacioni per librat eshte i centralizuar dhe i lehte per t'u lexuar."),
            ("Vendimmarrje me e mire", "Komentet dhe vleresimet ndihmojne klientet para blerjes."),
            ("Organizim me i mire", "Dyqani menaxhon libra, autore, kategori dhe porosi ne nje platforme te vetme."),
            ("Rritje e shitjeve", "Prezenca online e ben oferten me te dukshme per me shume kliente."),
        ],
        [2600, 6760],
    )

    heading(doc, "9. Veçori qe mund te shtohen ne te ardhmen")
    bullets(
        doc,
        [
            "Kerkim i avancuar sipas titullit, autorit, cmimit dhe kategorise.",
            "Shporte blerjeje me sasi te ndryshme per secilin liber.",
            "Pagesa online me kartele bankare ose metoda te tjera digjitale.",
            "Njoftime per statusin e porosise permes email-it.",
            "Panel statistikor per shitjet, librat me te kerkuar dhe vleresimet.",
            "Rekomandime te personalizuara bazuar ne interesat e perdoruesit.",
        ],
    )

    heading(doc, "10. Perfundim")
    para(
        doc,
        "Platforma e-commerce per shitjen e librave ofron nje menyre moderne per prezantimin dhe shitjen "
        "e librave online. Ajo bashkon katalogun, llogarine e perdoruesit, favorite, komente, vleresime "
        "dhe porosi ne nje sistem te vetem. Me zgjerime te metejshme si pagesa online, kerkimi i avancuar "
        "dhe rekomandimet personale, platforma mund te kthehet ne nje zgjidhje te plote per nje librari digjitale.",
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
